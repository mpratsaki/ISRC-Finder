"""
utils/docx_engine.py

Fully dynamic, template-faithful DOCX renderer for canonical LabelCopyData.
Supports nested tables, empty cells, and percentage-based performer logic.
"""

from __future__ import annotations

import io
import re
import unicodedata
from collections import OrderedDict
from collections.abc import Iterable, Mapping, Sequence
from copy import deepcopy
from datetime import datetime
from typing import Any

from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Twips
from docx.text.paragraph import Paragraph

# Word's own default table-cell inset (left/right) when a cell defines no
# explicit <w:tcMar>. Used only as a last-resort fallback when we have to
# compute a tab-stop position ourselves (see _ensure_right_tab_stop).
_DEFAULT_CELL_MARGIN_TWIPS = 108

class DocxTemplateError(ValueError):
    """Raised when the private template no longer matches the expected layout."""

DOCX_FIXED_LABEL_ORDER = (
    "Composer(s)",
    "Author(s)",
    "Producer(s)",
    "Recording Engineer(s)",
    "Mixing Engineer(s)",
    "Mastering Engineer(s)",
    "Vocalist(s)",
    "Rapper(s)",
)

# Canonical role IDs (from utils/label_copy_engine.ROLE_DEFINITIONS) that are
# instrumentalist credits rather than lead performers: they're listed after
# vocalists/rappers, as "Name (Instrument)", with no percentage split. Add
# new instrument role IDs here if label_copy_engine grows more of them.
INSTRUMENTALIST_ROLE_IDS = (
    "guitarist", 
    "bassist", 
    "drummer", 
    "keyboardist", 
    "programmer", 
    "performer"
)

# The private template historically had a single row reserved for guitar
# credits only, carrying a known typo ("Guirtarist(s):" - see
# label_copy_engine.py). We now reuse that same row (whichever spelling the
# template actually has) to list ALL instrumentalists, not just guitarists.
INSTRUMENTALIST_ROW_LABELS = ("Guitarist(s):", "Guirtarist(s):")

def _clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())

def _unique_texts(values: Iterable[Any]) -> list[str]:
    output: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = _clean_text(value)
        key = text.casefold()
        if not text or key in seen:
            continue
        seen.add(key)
        output.append(text)
    return output

def _as_int(value: Any, default: int = 0) -> int:
    if value is None or isinstance(value, bool):
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

def _duration_from_legacy_text(value: Any) -> int | None:
    text = _clean_text(value)
    if not text:
        return None
    parts = text.split(":")
    if not all(part.isdigit() for part in parts):
        return None
    numbers = [int(part) for part in parts]
    if len(numbers) == 2:
        minutes, seconds = numbers
        return ((minutes * 60) + seconds) * 1000
    if len(numbers) == 3:
        hours, minutes, seconds = numbers
        return ((hours * 3600) + (minutes * 60) + seconds) * 1000
    return None

def _duration_ms(entity: Mapping[str, Any]) -> int | None:
    for canonical_key in ("duration_ms", "total_duration_ms"):
        if entity.get(canonical_key) is not None:
            return max(_as_int(entity.get(canonical_key)), 0)
    for legacy_key in ("duration", "total_duration"):
        parsed = _duration_from_legacy_text(entity.get(legacy_key))
        if parsed is not None:
            return parsed
    return None

def format_duration_docx(duration_ms: Any) -> str:
    if duration_ms is None:
        return ""
    total_seconds = max(_as_int(duration_ms), 0) // 1000
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

def _is_simple_text_run(run) -> bool:
    allowed_tags = {qn("w:rPr"), qn("w:t")}
    return all(child.tag in allowed_tags for child in run._r)

def _text_runs_with_offsets(paragraph: Paragraph):
    runs = list(paragraph.runs)
    offsets = []
    cursor = 0
    for run in runs:
        text = run.text or ""
        start = cursor
        cursor += len(text)
        offsets.append((run, start, cursor))
    return runs, offsets, cursor

def _replace_range_across_runs(paragraph: Paragraph, start: int, end: int, replacement: str) -> None:
    runs, offsets, total_length = _text_runs_with_offsets(paragraph)
    if start < 0 or end < start or end > total_length:
        raise DocxTemplateError("Invalid run-level replacement range.")

    if not runs:
        paragraph.add_run(replacement)
        return

    if start == end == total_length:
        target = next((run for run in reversed(runs) if _is_simple_text_run(run)), runs[-1])
        target.text = f"{target.text}{replacement}"
        return

    first_index = None
    last_index = None
    for index, (_, run_start, run_end) in enumerate(offsets):
        if first_index is None and run_start <= start < run_end:
            first_index = index
        if run_start < end <= run_end:
            last_index = index
            break

    if first_index is None:
        for index, (_, run_start, _) in enumerate(offsets):
            if run_start == start:
                first_index = index
                break
    if last_index is None and end == total_length:
        last_index = len(offsets) - 1

    if first_index is None or last_index is None:
        raise DocxTemplateError("Could not map template text to its runs.")

    first_run, first_start, _ = offsets[first_index]
    last_run, last_start, _ = offsets[last_index]
    first_text = first_run.text or ""
    last_text = last_run.text or ""
    prefix = first_text[: max(start - first_start, 0)]
    suffix = last_text[max(end - last_start, 0) :]

    if first_index == last_index:
        first_run.text = f"{prefix}{replacement}{suffix}"
        return

    first_run.text = f"{prefix}{replacement}"
    for index in range(first_index + 1, last_index):
        offsets[index][0].text = ""
    last_run.text = suffix

def _replace_first_across_runs(paragraph: Paragraph, old: str, new: str, *, required: bool = True) -> bool:
    full_text = "".join(run.text or "" for run in paragraph.runs)
    start = full_text.find(old)
    if start < 0:
        if required:
            raise DocxTemplateError(f"The template paragraph does not contain the expected token: {old!r}")
        return False
    _replace_range_across_runs(paragraph, start, start + len(old), new)
    return True

def _set_label_value(paragraph: Paragraph, label: str, value: Any) -> None:
    full_text = "".join(run.text or "" for run in paragraph.runs)
    label_start = full_text.find(label)
    if label_start < 0:
        raise DocxTemplateError(f"The template paragraph does not contain the expected label: {label!r}")

    label_end = label_start + len(label)
    clean_value = str(value or "").strip()
    replacement = f" {clean_value}" if clean_value else " "
    _replace_range_across_runs(paragraph, label_end, len(full_text), replacement)

def _ensure_table_row_exists(paragraphs: list[Paragraph], anchor_label: str, new_label: str) -> None:
    """
    Δυναμική εισαγωγή νέας γραμμής στον πίνακα (κλωνοποιώντας μια υπάρχουσα) 
    αν απουσιάζει το new_label από το template.
    """
    for p in paragraphs:
        if new_label in p.text:
            return

    for p in paragraphs:
        if anchor_label in p.text:
            tc = p._element.getparent()
            while tc is not None and tc.tag != qn('w:tc'):
                tc = tc.getparent()
            
            if tc is not None and tc.tag == qn('w:tc'):
                tr = tc.getparent()
                if tr is not None and tr.tag == qn('w:tr'):
                    new_tr = deepcopy(tr)
                    for attribute in (qn("w14:paraId"), qn("w14:textId")):
                        new_tr.attrib.pop(attribute, None)
                    
                    tr.addnext(new_tr)
                    
                    new_ps = [Paragraph(node, p._parent) for node in new_tr.xpath('.//w:p')]
                    paragraphs.extend(new_ps)
                    
                    for np in new_ps:
                        if anchor_label in np.text:
                            _replace_first_across_runs(np, anchor_label, new_label, required=False)
                    
                    tc_elements = list(new_tr.findall(qn('w:tc')))
                    try:
                        old_tc_elements = list(tr.findall(qn('w:tc')))
                        tc_idx = old_tc_elements.index(tc)
                        
                        if tc_idx + 1 < len(tc_elements):
                            next_tc = tc_elements[tc_idx + 1]
                            next_tc_ps = list(next_tc.xpath('.//w:p'))
                            if next_tc_ps:
                                target_p = Paragraph(next_tc_ps[0], p._parent)
                                target_p.clear()
                                for extra_p_node in next_tc_ps[1:]:
                                    next_tc.remove(extra_p_node)
                    except ValueError:
                        pass
                    return

def _set_dynamic_label_value(paragraphs: list[Paragraph], label: str, value: Any) -> None:
    """Smart lookup supporting proper line breaks inside DOCX table cells without losing format."""
    raw_val = str(value or "").strip()
    for i, p in enumerate(paragraphs):
        if label in p.text:
            if len(p.text.strip()) > len(label.strip()) + 1 and not p.text.strip().endswith(":"):
                _set_label_value(p, label, raw_val)
            else:
                target_p = None
                tc = p._element.getparent()
                while tc is not None and tc.tag != qn('w:tc'):
                    tc = tc.getparent()
                    
                if tc is not None and tc.tag == qn('w:tc'):
                    tr = tc.getparent()
                    if tr is not None and tr.tag == qn('w:tr'):
                        tc_elements = list(tr.findall(qn('w:tc')))
                        try:
                            tc_index = tc_elements.index(tc)
                            if tc_index + 1 < len(tc_elements):
                                next_tc = tc_elements[tc_index + 1]
                                next_tc_ps = list(next_tc.iter(qn('w:p')))
                                if next_tc_ps:
                                    target_p = Paragraph(next_tc_ps[0], p._parent)
                                    for extra_p_node in next_tc_ps[1:]:
                                        parent_node = extra_p_node.getparent()
                                        if parent_node is not None:
                                            parent_node.remove(extra_p_node)
                        except ValueError:
                            pass
                
                if target_p is None and i + 1 < len(paragraphs):
                    target_p = paragraphs[i+1]
                    
                if target_p is not None:
                    lines = raw_val.split('\n')
                    if len(lines) == 1:
                        _replace_range_across_runs(target_p, 0, len(target_p.text), raw_val)
                    else:
                        rPr = deepcopy(target_p.runs[0]._r.rPr) if target_p.runs and target_p.runs[0]._r.rPr is not None else None
                        target_p.clear()
                        for idx, line in enumerate(lines):
                            run = target_p.add_run(line)
                            if rPr is not None:
                                run._r.append(deepcopy(rPr))
                            if idx < len(lines) - 1:
                                run.add_break()
            return

def _format_performers_with_percentages(names: list[str], total_performers: int) -> str:
    """Splits 100% ownership among multiple performers via newlines."""
    if not names or total_performers == 0:
        return ""
    percentage = 100.0 / total_performers
    pct_str = f"{int(percentage)}%" if percentage.is_integer() else f"{percentage:.2f}%"
    return "\n".join(f"{name}\t\t{pct_str}" for name in names)

def _format_rights_value(line: Any, fallback_owner: str = "") -> str:
    if not isinstance(line, Mapping):
        return ""
    year = _as_int(line.get("year"), 0)
    owner = _clean_text(line.get("owner")) or _clean_text(fallback_owner)
    parts = [str(year)] if year else []
    if owner:
        parts.append(owner)
    return " ".join(parts)

def _remove_clone_identity_attributes(element) -> None:
    for attribute in (qn("w14:paraId"), qn("w14:textId")):
        element.attrib.pop(attribute, None)

def _clone_track_blocks(document: _DocumentType, track_count: int) -> list[list[Any]]:
    body = document.element.body
    start_elem = None
    end_elem = None
    
    for elem in body:
        text = "".join(node.text for node in elem.iter(qn('w:t')) if node.text)
        if "Track 1" in text and start_elem is None:
            start_elem = elem
        if start_elem is not None and ("Rapper(s):" in text or "Guirtarist(s):" in text or "Guitarist(s):" in text):
            end_elem = elem
            break
            
    if start_elem is None or end_elem is None:
        raise DocxTemplateError("Αδυναμία δυναμικού εντοπισμού του Track block. Βεβαιωθείτε ότι το template περιέχει 'Track 1' και τελειώνει με τον τελευταίο performer (π.χ. 'Rapper(s):').")
        
    started = False
    original_nodes = []
    for elem in body:
        if elem is start_elem:
            started = True
        if started:
            original_nodes.append(elem)
        if elem is end_elem:
            break
            
    # Each track lives in its own standalone <w:tbl>, so consecutive tracks
    # sit flush against one another with nothing between them (Word doesn't
    # add any gap between adjacent tables on its own). The template already
    # defines exactly this kind of spacer - a small empty paragraph with
    # `w:spacing w:after="120"` - immediately above the very first track
    # block, to separate it from the "TRACKLIST DETAILS" header. We reuse
    # that same paragraph as a template and clone it between every pair of
    # generated track blocks, so the last row of one track ("PERFORMED BY" /
    # Rapper(s)) never touches the next track's title row.
    spacer_template = original_nodes[0].getprevious()
    if spacer_template is None or spacer_template.tag != qn('w:p'):
        spacer_template = None

    insertion_index = body.index(original_nodes[0])
    blocks = []
    
    for track_index in range(track_count):
        block_nodes = []
        for original_node in original_nodes:
            clone = deepcopy(original_node)
            _remove_clone_identity_attributes(clone)
            body.insert(insertion_index, clone)
            insertion_index += 1
            block_nodes.append(clone)
        blocks.append(block_nodes)

        if track_index < track_count - 1:
            if spacer_template is not None:
                spacer_clone = deepcopy(spacer_template)
            else:
                spacer_clone = parse_xml(
                    f'<w:p {nsdecls("w")}><w:pPr>'
                    f'<w:spacing w:after="120" w:lineRule="auto"/>'
                    f'</w:pPr></w:p>'
                )
            _remove_clone_identity_attributes(spacer_clone)
            body.insert(insertion_index, spacer_clone)
            insertion_index += 1
        
    for original_node in original_nodes:
        body.remove(original_node)
        
    section_properties = body.sectPr
    if section_properties is not None and body[-1] is not section_properties:
        body.remove(section_properties)
        body.append(section_properties)
        
    return blocks

def _credit_names(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [_clean_text(value)] if _clean_text(value) else []
    if isinstance(value, Mapping):
        if "names" in value:
            return _credit_names(value.get("names"))
        name = value.get("name")
        return [_clean_text(name)] if _clean_text(name) else []
    if isinstance(value, Iterable):
        names: list[str] = []
        for item in value:
            names.extend(_credit_names(item))
        return _unique_texts(names)
    return [_clean_text(value)] if _clean_text(value) else []

def _canonical_credits_for_render(track: Mapping[str, Any]) -> tuple[dict[str, list[str]], dict[str, str]]:
    from utils.label_copy_engine import ROLE_DEFINITIONS, resolve_canonical_roles
    raw_credits = track.get("credits")
    raw_labels = track.get("credit_labels")
    if not isinstance(raw_credits, Mapping):
        return {}, {}

    credits: dict[str, list[str]] = {}
    labels: dict[str, str] = {}
    for raw_role, raw_names in raw_credits.items():
        names = _credit_names(raw_names)
        if not names:
            continue

        role_text = str(raw_role)
        if role_text in ROLE_DEFINITIONS or role_text.startswith("other:"):
            resolved = [(role_text, _clean_text(raw_labels.get(role_text)) if isinstance(raw_labels, Mapping) else "")]
        else:
            resolved = resolve_canonical_roles(role_text)

        for role_id, display_label in resolved:
            if role_id in ROLE_DEFINITIONS and not display_label:
                display_label = ROLE_DEFINITIONS[role_id]["display_label"]
            if role_id.startswith("other:") and not display_label:
                display_label = role_text
            credits[role_id] = _unique_texts([*credits.get(role_id, []), *names])
            labels.setdefault(role_id, display_label or role_text)

    ordered: dict[str, list[str]] = OrderedDict()
    for role_id in ROLE_DEFINITIONS:
        if credits.get(role_id):
            ordered[role_id] = credits[role_id]
    for role_id, names in credits.items():
        if role_id not in ordered:
            ordered[role_id] = names
    return dict(ordered), labels

def _docx_credit_values(track: Mapping[str, Any]) -> tuple[dict[str, list[str]], list[str], list[str]]:
    from utils.label_copy_engine import ROLE_DEFINITIONS
    credits, labels = _canonical_credits_for_render(track)
    fixed: dict[str, list[str]] = {label: [] for label in DOCX_FIXED_LABEL_ORDER}
    other_lines: list[str] = []
    instrumentalist_lines: list[str] = []

    for role_id, names in credits.items():
        is_instrument = role_id in INSTRUMENTALIST_ROLE_IDS
        
        # Πιάνουμε δυναμικά Tidal 'other:' roles που είναι όργανα (π.χ. other:acoustic_guitar, other:violin)
        if role_id.startswith("other:"):
            role_lower = role_id.lower()
            if any(i in role_lower for i in (
                "violin", "cello", "viola", "string", "trumpet", 
                "sax", "horn", "flute", "brass", "woodwind", 
                "synth", "piano", "organ", "instrument", "percussion"
            )):
                is_instrument = True

        if is_instrument:
            definition = ROLE_DEFINITIONS.get(role_id)
            instrument_label = _clean_text(labels.get(role_id))
            if not instrument_label:
                instrument_label = _clean_text(
                    (definition or {}).get("display_label") or role_id.title()
                )
            # Αποφυγή άβολων συντακτικών όπως "John Doe (Performed by)"
            if instrument_label.lower() == "performed by":
                instrument_label = "Performer"
                
            instrumentalist_lines.extend(f"{name} ({instrument_label})" for name in names)
            continue

        definition = ROLE_DEFINITIONS.get(role_id)
        docx_label = definition.get("docx_label") if definition else None
        if docx_label in fixed:
            fixed[docx_label] = _unique_texts([*fixed[docx_label], *names])
            continue

        display_label = labels.get(role_id)
        if not display_label and definition:
            display_label = definition.get("display_label") or definition.get("pdf_label")
        display_label = _clean_text(display_label) or role_id.removeprefix("other:").replace("_", " ").title()
        display_label = display_label.rstrip(":")
        other_lines.append(f"{display_label}: {', '.join(names)}")

    return fixed, other_lines, instrumentalist_lines

def _fill_release_header(paragraphs: list[Paragraph], data: Mapping[str, Any]) -> None:
    for p in paragraphs:
        if "PROJECT NAME" in p.text:
            _replace_first_across_runs(p, "PROJECT NAME", _clean_text(data.get("project_name")), required=False)
        elif "Project Name" in p.text:
            _replace_first_across_runs(p, "Project Name", _clean_text(data.get("project_name")), required=False)
        if "DD/MM/YYYY" in p.text:
            _replace_first_across_runs(p, "DD/MM/YYYY", _clean_text(data.get("issue_date")), required=False)
            
    _set_dynamic_label_value(paragraphs, "Artist(s):", ", ".join(_unique_texts(data.get("artists", []))))
    _set_dynamic_label_value(paragraphs, "Product Type:", data.get("product_type"))
    _set_dynamic_label_value(paragraphs, "UPC:", data.get("upc"))
    _set_dynamic_label_value(paragraphs, "Release Date:", data.get("release_date"))
    _set_dynamic_label_value(paragraphs, "Label Imprint:", data.get("label_imprint"))
    _set_dynamic_label_value(paragraphs, "Company:", data.get("company"))
    _set_dynamic_label_value(paragraphs, "Metadata Language:", data.get("metadata_language") or data.get("metadata_language_suggestion"))
    _set_dynamic_label_value(paragraphs, "Genre / Subgenre:", f"{data.get('genre')} / {data.get('subgenre')}")
    _set_dynamic_label_value(paragraphs, "Total Duration:", format_duration_docx(_duration_ms(data)))

def _apply_rPr(run, source_rPr) -> None:
    """Clone `source_rPr`'s formatting onto `run`, preserving valid OOXML ordering.

    `w:rPr` must be the first child of `w:r` per the schema. Using
    `get_or_add_rPr()` guarantees it is inserted (or already sits) in the
    correct position; we then swap in the source's formatting children
    rather than raw-appending a second `w:rPr`, which is what produced the
    invalid `<w:r><w:t>...</w:t><w:rPr>...</w:rPr></w:r>` ordering before.
    """
    if source_rPr is None:
        return
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        rPr.remove(child)
    for child in deepcopy(source_rPr):
        rPr.append(child)

def _append_styled_run(paragraph: Paragraph, text: str, source_rPr) -> None:
    run = paragraph.add_run(text)
    _apply_rPr(run, source_rPr)

def _grid_col_widths_twips(tbl_element) -> list[int]:
    widths = []
    tblGrid = tbl_element.find(qn('w:tblGrid'))
    if tblGrid is None:
        return widths
    for gridCol in tblGrid.findall(qn('w:gridCol')):
        try:
            widths.append(int(float(gridCol.get(qn('w:w'), 0))))
        except (TypeError, ValueError):
            widths.append(0)
    return widths

def _cell_margin_twips(tcPr, tblPr, side: str) -> int:
    """Resolve a cell's left/right inset: explicit w:tcMar, else the
    table-wide w:tblCellMar default, else Word's own built-in default."""
    tag = qn(f'w:{side}')
    for pr in (tcPr, tblPr):
        if pr is None:
            continue
        mar = pr.find(qn('w:tcMar')) if pr.tag == qn('w:tcPr') else pr.find(qn('w:tblCellMar'))
        if mar is None:
            continue
        node = mar.find(tag)
        if node is not None and node.get(qn('w:w')) is not None:
            try:
                return int(float(node.get(qn('w:w'))))
            except (TypeError, ValueError):
                continue
    return _DEFAULT_CELL_MARGIN_TWIPS

def _measure_cell_inner_width_twips(tc_element) -> int | None:
    """Best-effort inner (content) width of the table cell containing
    `tc_element`, in twips: explicit w:tcW if set, otherwise the sum of the
    w:tblGrid columns the cell spans, minus its left/right margins.
    Returns None if it can't be determined (caller should fall back)."""
    tcPr = tc_element.find(qn('w:tcPr'))

    tr = tc_element.getparent()
    if tr is None or tr.tag != qn('w:tr'):
        return None
    tbl = tr.getparent()
    if tbl is None or tbl.tag != qn('w:tbl'):
        return None
    tblPr = tbl.find(qn('w:tblPr'))

    left_margin = _cell_margin_twips(tcPr, tblPr, "left")
    right_margin = _cell_margin_twips(tcPr, tblPr, "right")

    explicit_width = None
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None and tcW.get(qn('w:type')) == 'dxa':
            try:
                explicit_width = int(float(tcW.get(qn('w:w'), 0)))
            except (TypeError, ValueError):
                explicit_width = None

    if explicit_width:
        return max(explicit_width - left_margin - right_margin, 0)

    # No explicit cell width: derive it from the table's column grid,
    # matching however many columns this cell spans starting at its
    # position among its row's cells.
    grid_widths = _grid_col_widths_twips(tbl)
    if not grid_widths:
        return None

    tc_siblings = list(tr.findall(qn('w:tc')))
    try:
        tc_index = tc_siblings.index(tc_element)
    except ValueError:
        return None

    col_cursor = 0
    for sibling in tc_siblings[:tc_index]:
        sibling_pr = sibling.find(qn('w:tcPr'))
        span_node = sibling_pr.find(qn('w:gridSpan')) if sibling_pr is not None else None
        span = int(span_node.get(qn('w:val'))) if span_node is not None else 1
        col_cursor += span

    span_node = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
    own_span = int(span_node.get(qn('w:val'))) if span_node is not None else 1

    spanned = grid_widths[col_cursor: col_cursor + own_span]
    if not spanned:
        return None

    return max(sum(spanned) - left_margin - right_margin, 0)

def _ensure_right_tab_stop(paragraph: Paragraph) -> None:
    """Guarantee the paragraph has a right-aligned tab stop at the inner
    right edge of its containing table cell.

    The Stay Independent template already ships this tab stop on the track
    title row (so this is normally a no-op), but computing it defensively
    means the layout keeps working even if the template is edited by hand
    and the tab stop is dropped.
    """
    existing = list(paragraph.paragraph_format.tab_stops)
    if any(stop.alignment == WD_TAB_ALIGNMENT.RIGHT for stop in existing):
        return

    width_twips = _measure_cell_inner_width_twips(paragraph._p.getparent())
    if width_twips is None or width_twips <= 0:
        return

    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Twips(width_twips), WD_TAB_ALIGNMENT.RIGHT
    )

def _fill_track_block(block_nodes: list[Any], track: Mapping[str, Any], display_number: int) -> None:
    paragraphs = []
    for node in block_nodes:
        paragraphs.extend(Paragraph(p, None) for p in node.xpath('.//w:p'))

    for p in paragraphs:
        if "Track 1" in p.text and "Duration:" in p.text:
            title = _clean_text(track.get("title"))
            duration_str = format_duration_docx(_duration_ms(track))

            runs = p.runs
            left_rPr = deepcopy(runs[0]._r.rPr) if runs and runs[0]._r.rPr is not None else None
            right_rPr = (
                deepcopy(runs[1]._r.rPr)
                if len(runs) > 1 and runs[1]._r.rPr is not None
                else deepcopy(left_rPr)
            )

            left_text = f"Track {display_number}: {title}"
            right_text = f"Duration: {duration_str}"

            p.clear()

            _append_styled_run(p, left_text, left_rPr)
            _append_styled_run(p, "\t", right_rPr)
            _append_styled_run(p, right_text, right_rPr)

            _ensure_right_tab_stop(p)
            break
            
    _set_dynamic_label_value(paragraphs, "Primary Artist(s):", ", ".join(_unique_texts(track.get("primary_artists", []))))
    _set_dynamic_label_value(paragraphs, "Featured Artist(s):", ", ".join(_unique_texts(track.get("featured_artists", []))))
    _set_dynamic_label_value(paragraphs, "ISRC:", track.get("isrc"))
    _set_dynamic_label_value(paragraphs, "Genre:", track.get("genre"))
    _set_dynamic_label_value(paragraphs, "Lyrics Language:", track.get("lyrics_language") or track.get("lyrics_language_suggestion"))
    _set_dynamic_label_value(paragraphs, "Parental Advisory:", track.get("parental_advisory"))
    
    # Λαμβάνουμε και τις τρεις λίστες χωρίς error
    fixed_credits, other_lines, instrumentalist_lines = _docx_credit_values(track)
    
    for label, names in fixed_credits.items():
        if label not in ("Vocalist(s)", "Rapper(s)"):
            _set_dynamic_label_value(paragraphs, f"{label}:", ", ".join(names))
        
    vocalists = fixed_credits.get("Vocalist(s)", [])
    rappers = fixed_credits.get("Rapper(s)", [])
    
    total_performers = len(vocalists) + len(rappers)
    
    _set_dynamic_label_value(paragraphs, "Vocalist(s):", _format_performers_with_percentages(vocalists, total_performers))
    _set_dynamic_label_value(paragraphs, "Rapper(s):", _format_performers_with_percentages(rappers, total_performers))
    
    # 1. Αν υπάρχει κάποιο legacy label τύπου Guitarist(s):, το μετονομάζουμε.
    for p in paragraphs:
        for old_label in INSTRUMENTALIST_ROW_LABELS:
            if old_label in p.text:
                _replace_first_across_runs(p, old_label, "Musician(s):", required=False)
                break

    # 2. Εισαγωγή γραμμής Musician(s): αν δεν υπάρχει (αγκυρωμένη στους Rapper)
    if instrumentalist_lines:
        _ensure_table_row_exists(paragraphs, "Rapper(s):", "Musician(s):")
        _set_dynamic_label_value(paragraphs, "Musician(s):", "\n".join(instrumentalist_lines))
    else:
        _set_dynamic_label_value(paragraphs, "Musician(s):", "")

    # 3. Εισαγωγή γραμμής Other Credits: αν δεν υπάρχει 
    if other_lines:
        # Ψάχνει πρώτα να "κουμπώσει" κάτω από τους Musician(s), αλλιώς κάτω από τους Rapper(s)
        _ensure_table_row_exists(paragraphs, "Musician(s):", "Other Credits:")
        _ensure_table_row_exists(paragraphs, "Rapper(s):", "Other Credits:")
        _set_dynamic_label_value(paragraphs, "Other Credits:", "\n".join(other_lines))
    else:
        _set_dynamic_label_value(paragraphs, "Other Credits:", "")
def generate_label_copy_docx(template_bytes: bytes, data: Mapping[str, Any]) -> io.BytesIO:
    if not isinstance(template_bytes, (bytes, bytearray)) or not template_bytes:
        raise ValueError("Το Label Copy template είναι κενό ή μη έγκυρο.")
    if not isinstance(data, Mapping):
        raise TypeError("data must be a LabelCopyData mapping")

    try:
        document = Document(io.BytesIO(bytes(template_bytes)))
    except Exception as exc:
        raise DocxTemplateError("Αδυναμία ανοίγματος του Label Copy DOCX template.") from exc

    tracks = data.get("tracks")
    if not isinstance(tracks, list) or not tracks:
        raise ValueError("Το LabelCopyData δεν περιέχει tracks.")

    track_blocks = _clone_track_blocks(document, len(tracks))
    
    header_nodes = []
    for elem in document.element.body:
        text = "".join(node.text for node in elem.iter(qn('w:t')) if node.text)
        if "Track 1" in text:
            break
        header_nodes.append(elem)
        
    header_paragraphs = [Paragraph(p, None) for node in header_nodes for p in node.xpath('.//w:p')]
    _fill_release_header(header_paragraphs, data)
    
    for display_number, (block, track) in enumerate(zip(track_blocks, tracks), start=1):
        if not isinstance(track, Mapping):
            continue
        _fill_track_block(block, track, display_number)
        
    footer_nodes = []
    for elem in reversed(document.element.body):
        text = "".join(node.text for node in elem.iter(qn('w:t')) if node.text)
        if "Rapper(s):" in text or "Guirtarist(s):" in text or "Guitarist(s):" in text:
            break
        footer_nodes.append(elem)
        
    footer_paragraphs = [Paragraph(p, None) for node in footer_nodes for p in node.xpath('.//w:p')]
    
    p_text = _format_rights_value(data.get("p_line"), data.get("label_imprint", ""))
    c_text = _format_rights_value(data.get("c_line"), data.get("label_imprint", ""))
    
    for p in footer_paragraphs:
        if "©" in p.text or "(C)" in p.text:
            _replace_range_across_runs(p, 0, len(p.text), f"© {c_text}")
        elif "℗" in p.text or "(P)" in p.text:
            _replace_range_across_runs(p, 0, len(p.text), f"℗ {p_text}")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output

def make_label_copy_filename(project_name: Any, *, extension: str = "docx", issue_date: Any = None) -> str:
    normalized = unicodedata.normalize("NFKD", _clean_text(project_name))
    ascii_title = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_title = re.sub(r"[^A-Za-z0-9]+", "_", ascii_title).strip("_") or "Release"
    date_token = ""
    issue_text = _clean_text(issue_date)
    for date_format in ("%d/%m/%Y", "%Y-%m-%d", "%Y%m%d"):
        try:
            date_token = datetime.strptime(issue_text, date_format).strftime("%Y%m%d")
            break
        except ValueError:
            continue
    if not date_token:
        date_token = datetime.now().strftime("%Y%m%d")
    clean_extension = re.sub(r"[^A-Za-z0-9]", "", str(extension or "docx")).lower() or "docx"
    return f"LabelCopy_{ascii_title}_{date_token}.{clean_extension}"

render_label_copy_docx = generate_label_copy_docx

__all__ = [
    "DocxTemplateError",
    "format_duration_docx",
    "generate_label_copy_docx",
    "make_label_copy_filename",
    "render_label_copy_docx",
]
