"""Round-trip regression test for utils.docx_engine."""

from __future__ import annotations

import io
from docx import Document
from docx.oxml.ns import qn

from utils.docx_engine import generate_label_copy_docx


TRACK_BLOCK = [
    "Track 1                                                                                          Duration: 00:00:00",
    " Primary Artist(s): ",
    "Featured Artist(s): ",
    "ISRC:",
    "Genre:",
    "Subgenre:",
    "Lyrics Language:",
    "Parental Advisory:",
    "Written By",
    "Composer(s):",
    "Author(s):",
    "Produced by",
    "Producer(s):",
    "Recording Engineer(s):",
    "Mixing Engineer(s):",
    "Mastering Engineer(s):",
    "Performed By",
    "Vocalist(s):",
    "Guirtarist(s):",
    "Other Credits:",
]


def _make_template_bytes() -> bytes:
    document = Document()
    first = document.add_paragraph()

    # Split both placeholders across runs to exercise cross-run mutation.
    run = first.add_run("             Label Copy: Proj")
    run.bold = True
    run = first.add_run("ect Name                      Issue Date: DD/")
    run.bold = True
    run = first.add_run("MM/YYYY ")
    run.bold = True

    release_lines = [
        " ",
        "Artist(s): ",
        "Product Type: ",
        "UPC: ",
        "Release Date: ",
        "Label Imprint: ",
        "Company: Stay Independent",
        "Metadata Language: Greek (GR)",
        "Genre: ",
        "Subgenre: ",
        "Total Duration: 1:17:36",
        "————————————————————————————————————————",
    ]
    for text in release_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True

    for index, text in enumerate(TRACK_BLOCK):
        paragraph = document.add_paragraph()
        if index == 0:
            run = paragraph.add_run("Track 1")
            run.bold = True
            for fragment in (
                "                                                                                          Duration: 00:0",
                "0",
                ":",
                "00",
            ):
                run = paragraph.add_run(fragment)
                run.bold = True
        elif text == "Guirtarist(s):":
            run = paragraph.add_run("Guirtarist")
            run.bold = True
            run = paragraph.add_run("(s):")
            run.bold = True
        else:
            run = paragraph.add_run(text)
            run.bold = True
            if text in {"Written By", "Produced by", "Performed By"}:
                run.underline = True

    paragraph = document.add_paragraph()
    paragraph.add_run("—————————————————————————————————————————(P) ").bold = True
    paragraph = document.add_paragraph()
    paragraph.add_run("(C) ").bold = True
    document.add_paragraph("")

    assert len(document.paragraphs) == 36
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _fixture_data() -> dict:
    base_track = {
        "disc_number": 1,
        "primary_artists": ["Main Artist"],
        "featured_artists": [],
        "genre": "Alternative",
        "subgenre": "Indie Pop",
        "lyrics_language": "Greek (GR)",
        "parental_advisory": "Non-Applicable",
        "credits": {
            "composer": ["Composer Name"],
            "lyricist": ["Lyricist Name"],
            "producer": ["Producer Name"],
            "recording_engineer": ["Recording Engineer"],
            "mixing_engineer": ["Mix Engineer"],
            "mastering_engineer": ["Mastering Engineer"],
            "vocalist": ["Main Artist"],
            "guitarist": ["Guitar Player"],
            "arranger": ["Arranger Name"],
        },
        "credit_labels": {"arranger": "Arranged by"},
    }
    return {
        "project_name": "Three Track Release",
        "issue_date": "22/07/2026",
        "artists": ["Main Artist"],
        "product_type": "EP",
        "upc": "012345678901",
        "release_date": "07/2026",
        "release_date_precision": "month",
        "label_imprint": "Test Imprint",
        "company": "Stay Independent",
        "publisher": "Stay Independent",
        "metadata_language": "Greek (GR)",
        "genre": "Alternative",
        "subgenre": "Indie Pop",
        "total_duration_ms": 600_000,
        "p_line": {"year": 2026, "owner": "Test Imprint"},
        "c_line": {"year": 2026, "owner": "Stay Independent"},
        "tracks": [
            {
                **base_track,
                "number": 1,
                "track_number": 1,
                "title": "First Song",
                "duration_ms": 180_000,
                "isrc": "GRABC2600001",
            },
            {
                **base_track,
                "number": 2,
                "track_number": 2,
                "title": "Second Song",
                "duration_ms": 195_000,
                "isrc": "GRABC2600002",
                "featured_artists": ["Guest Artist"],
            },
            {
                **base_track,
                "number": 3,
                "track_number": 3,
                "title": "Third Song",
                "duration_ms": 225_000,
                "isrc": "GRABC2600003",
            },
        ],
        "warnings": [],
    }


def test_three_track_template_round_trip() -> None:
    output = generate_label_copy_docx(_make_template_bytes(), _fixture_data())
    assert output.getvalue().startswith(b"PK")

    rendered = Document(io.BytesIO(output.getvalue()))

    # Original 36 paragraphs - one 20-paragraph block + three cloned blocks.
    assert len(rendered.paragraphs) == 36 - 20 + (3 * 20)

    track_headers = [
        paragraph.text
        for paragraph in rendered.paragraphs
        if paragraph.text.startswith("Track ")
    ]
    assert len(track_headers) == 3
    assert track_headers[0].startswith("Track 1: First Song")
    assert track_headers[1].startswith("Track 2: Second Song")
    assert track_headers[2].startswith("Track 3: Third Song")

    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Project Name" not in all_text
    assert "DD/MM/YYYY" not in all_text
    assert "00:00:00" not in all_text
    assert "Three Track Release" in all_text
    assert "Total Duration: 00:10:00" in all_text
    assert "Guirtarist(s): Guitar Player" in all_text
    assert "Arranged by: Arranger Name" in all_text
    assert "(P) 2026 Test Imprint" in all_text
    assert "(C) 2026 Stay Independent" in all_text

    # The body-level section properties remain the final XML child.
    assert rendered.element.body[-1].tag == qn("w:sectPr")
